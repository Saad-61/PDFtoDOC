import time
import logging
import asyncio
import re
import math
from pathlib import Path
from typing import List, Optional, Callable, Dict, Any
import fitz
try:
    from pdf2docx import Converter
except ImportError:  # pragma: no cover
    Converter = Any  # type: ignore
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.core.config import settings
from app.core.exceptions import ConversionException, FileValidationException

logger = logging.getLogger("pdf2docx.engine")


def render_rich_footer_runs(
    text: str,
    default_font: str = "Calibri",
    default_sz: int = 16,
    default_color: str = "71717A",
) -> str:
    """
    Splits text into regular runs and hyperlinked runs (URLs and Emails) with native blue color and single underline.
    """
    pattern = r'(https?://[^\s,]+|www\.[^\s,]+|[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)'
    tokens = re.split(pattern, text)
    runs_xml = []
    for token in tokens:
        if not token:
            continue
        if re.match(r'^(https?://|www\.)', token, re.I) or ('@' in token and '.' in token):
            # Hyperlink run: Theme Blue (#0563C1), single underline
            runs_xml.append(
                f'  <w:r>\n'
                f'    <w:rPr>\n'
                f'      <w:rFonts w:ascii="{default_font}" w:hAnsi="{default_font}"/>\n'
                f'      <w:sz w:val="{default_sz}"/>\n'
                f'      <w:color w:val="0563C1"/>\n'
                f'      <w:u w:val="single"/>\n'
                f'    </w:rPr>\n'
                f'    <w:t xml:space="preserve">{token}</w:t>\n'
                f'  </w:r>'
            )
        else:
            # Regular text run
            runs_xml.append(
                f'  <w:r>\n'
                f'    <w:rPr>\n'
                f'      <w:rFonts w:ascii="{default_font}" w:hAnsi="{default_font}"/>\n'
                f'      <w:sz w:val="{default_sz}"/>\n'
                f'      <w:color w:val="{default_color}"/>\n'
                f'    </w:rPr>\n'
                f'    <w:t xml:space="preserve">{token}</w:t>\n'
                f'  </w:r>'
            )
    return "\n".join(runs_xml)


def analyze_pdf_layout_structure(pdf_path: Path) -> Dict[str, Any]:
    """
    Comprehensive pre-flight analyzer that scans all pages of a PDF to extract:
    1. Diagonal / rotated text watermarks (e.g. 'Watermark', 'Confidential', 'Draft').
    2. Image classification (background watermark emblem vs foreground photo/signature).
    3. Footer classification (page numbers vs contact info vs empty).
    4. Text signatures for cleaning misplaced body paragraphs.
    """
    try:
        doc_pdf = fitz.open(str(pdf_path))
        num_pages = len(doc_pdf)
        if num_pages == 0:
            return {
                "watermark": None,
                "has_page_numbers": False,
                "has_text_footer": False,
                "footer_rows": [],
                "body_cleaning_signatures": [],
                "background_image_xrefs": set(),
            }

        watermark_meta = None
        watermark_pages = []
        has_page_numbers = False
        footer_candidate_lines = []
        body_cleaning_signatures = set()
        background_image_xrefs = set()
        page_w = 595.0

        for page_idx, page in enumerate(doc_pdf):
            page_h = page.rect.height
            page_w = page.rect.width
            page_area = page_w * page_h

            footer_top_y = page_h * 0.915  # Bottom ~8.5% margin for footers
            d = page.get_text("dict")

            # 1. Image analysis on this page
            page_imgs = page.get_images()
            for img_info in page_imgs:
                xref = img_info[0]
                img_rects = page.get_image_rects(xref)
                for r in img_rects:
                    img_area = r.width * r.height
                    is_large_centered = (r.width > page_w * 0.45 and r.height > page_h * 0.35)
                    is_huge = (img_area >= page_area * 0.25)
                    if is_large_centered or is_huge:
                        background_image_xrefs.add(xref)

            page_footer_lines = []

            # 2. Text block analysis
            for block in d.get("blocks", []):
                if block.get("type") == 0:  # Text
                    for line in block.get("lines", []):
                        dx, dy = line.get("dir", (1, 0))
                        line_bbox = line.get("bbox", (0, 0, 0, 0))
                        spans = line.get("spans", [])
                        line_text = "".join(s.get("text", "") for s in spans).strip()

                        # A. Check for diagonal text watermark (|dy| > 0.08 and size >= 14)
                        if abs(dy) > 0.08:
                            for s in spans:
                                txt = s.get("text", "").strip()
                                if len(txt) > 1 and s.get("size", 0) >= 14:
                                    if not watermark_meta:
                                        c = s.get("color", 0)
                                        r_col = (c >> 16) & 0xFF
                                        g_col = (c >> 8) & 0xFF
                                        b_col = c & 0xFF
                                        hex_col = f"{r_col:02x}{g_col:02x}{b_col:02x}".upper() if c != 0 else "2562EB"
                                        font_name = s.get("font", "DejaVu Sans")
                                        if "-" in font_name:
                                            font_name = font_name.split("-")[0]
                                        # In PDF coordinates (y down), visual upward angle:
                                        visual_deg = math.degrees(math.atan2(-dy, dx))
                                        # In Word DrawingML, rot is clockwise: (360 - visual_deg) * 60000
                                        rot_val = int(round((360.0 - visual_deg) * 60000)) % 21600000
                                        watermark_meta = {
                                            "text": txt,
                                            "size": s.get("size", 42),
                                            "font": font_name,
                                            "color": hex_col,
                                            "rot": rot_val if rot_val != 0 else 19800000,
                                        }
                                    if page_idx not in watermark_pages:
                                        watermark_pages.append(page_idx)

                        # B. Check for footer zone content (y >= footer_top_y)
                        elif line_bbox[1] >= footer_top_y or line_bbox[3] >= footer_top_y:
                            if line_text:
                                # Standalone page number (e.g. '1', '2', 'Page 1 of 3')
                                if line_text.isdigit() or re.match(r"^page\s*\d+(\s*of\s*\d+)?$", line_text, re.I):
                                    has_page_numbers = True
                                    body_cleaning_signatures.add(line_text)
                                # Contact / Address text footer
                                elif any(k in line_text.lower() for k in ("tel:", "email:", "e-mail:", "web:", "www.", "usa:", "copyright", "©")):
                                    page_footer_lines.append({
                                        "bbox": line_bbox,
                                        "text": line_text,
                                        "spans": spans,
                                    })
                                    if len(line_text) >= 4:
                                        body_cleaning_signatures.add(line_text)
                                        sub_parts = re.split(r'\s{2,}|\t|(?=Tel:)|(?=E-mail:)|(?=Email:)|(?=Web:)|(?=USA:)', line_text)
                                        for sp in sub_parts:
                                            sp_clean = sp.strip()
                                            if len(sp_clean) >= 4:
                                                body_cleaning_signatures.add(sp_clean)
                                        # Also add URLs and emails explicitly
                                        for link_match in re.findall(r'(https?://[^\s,]+|www\.[^\s,]+|[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', line_text):
                                            body_cleaning_signatures.add(link_match)

            if page_footer_lines:
                footer_candidate_lines.append(page_footer_lines)

        doc_pdf.close()

        if watermark_meta:
            watermark_meta["pages_present"] = watermark_pages

        # Build structured footer rows for text footers
        structured_footer_rows = []
        if footer_candidate_lines:
            sample_lines = max(footer_candidate_lines, key=len)
            mid_x = page_w * 0.48

            for line_item in sample_lines:
                spans = line_item.get("spans", [])
                line_text = line_item.get("text", "")
                left_parts = []
                right_parts = []

                for s in spans:
                    stext = s.get("text", "").strip()
                    if not stext:
                        continue
                    s_bbox = s.get("bbox", (0, 0, 0, 0))
                    if s_bbox[0] < mid_x:
                        left_parts.append(stext)
                    else:
                        right_parts.append(stext)

                left_str = " ".join(left_parts).strip()
                right_str = " ".join(right_parts).strip()

                if not right_str and any(k in line_text for k in ("Tel:", "E-mail:", "Email:", "Phone:", "www.", "Web:")):
                    parts = re.split(r'\s{2,}|\t|(?=Tel:)|(?=E-mail:)|(?=Email:)|(?=Phone:)', line_text)
                    parts = [p.strip() for p in parts if p.strip()]
                    if len(parts) >= 2:
                        left_str = parts[0]
                        right_str = " ".join(parts[1:])
                    else:
                        left_str = line_text
                        right_str = ""
                elif not left_str and not right_str:
                    left_str = line_text
                    right_str = ""

                sample_span = spans[0] if spans else {}
                c = sample_span.get("color", 0)
                r_col = (c >> 16) & 0xFF
                g_col = (c >> 8) & 0xFF
                b_col = c & 0xFF
                hex_col = f"{r_col:02x}{g_col:02x}{b_col:02x}".upper() if c != 0 else "71717A"
                if hex_col in ("000000", "FFFFFF"):
                    hex_col = "71717A"

                font_name = sample_span.get("font", "Calibri")
                if "-" in font_name:
                    font_name = font_name.split("-")[0]
                if font_name.lower().startswith("dejavu"):
                    font_name = "Calibri"

                font_sz_pt = max(7.5, min(9.5, sample_span.get("size", 8.0)))

                structured_footer_rows.append({
                    "left_text": left_str,
                    "right_text": right_str,
                    "font": font_name,
                    "size_half_pt": int(round(font_sz_pt * 2)),
                    "color": hex_col,
                })

        return {
            "watermark": watermark_meta,
            "has_page_numbers": has_page_numbers,
            "has_text_footer": len(structured_footer_rows) > 0,
            "footer_rows": structured_footer_rows,
            "body_cleaning_signatures": list(body_cleaning_signatures),
            "background_image_xrefs": background_image_xrefs,
        }
    except Exception as e:
        logger.warning(f"PDF pre-flight layout analysis issue: {e}")
        return {
            "watermark": None,
            "has_page_numbers": False,
            "has_text_footer": False,
            "footer_rows": [],
            "body_cleaning_signatures": [],
            "background_image_xrefs": set(),
        }


def apply_word_post_processing(
    docx_path: Path,
    layout_info: Optional[Dict[str, Any]] = None,
) -> bool:
    """
    Applies semantic layout refinements to the generated Word DOCX based on pre-flight inspection:
    1. Sets default View to 'Print Layout' in word/settings.xml so Word opens in standard layout.
    2. Injects modern DrawingML watermark into Header with correct counter-clockwise orientation matching PDF.
    3. Selectively sets behindDoc="1" for large background watermark emblems and behindDoc="0" for foreground photos/signatures.
    4. Cleans misplaced footer lines (including hyperlinked runs) and body page numbers from body paragraphs.
    5. Normalizes paragraph spacing so content fits cleanly on intended pages.
    6. Injects native page numbers or rich text footers (with styled hyperlinks) into Word footers.
    """
    try:
        doc = docx.Document(str(docx_path))

        # 1. Force Print Layout view mode by default in settings.xml
        try:
            settings_elm = doc.settings.element
            if not settings_elm.xpath("w:view"):
                view_xml = f'<w:view {nsdecls("w")} w:val="print"/>'
                settings_elm.insert(0, parse_xml(view_xml))
        except Exception as view_err:
            logger.warning(f"Failed to set print layout view setting: {view_err}")

        # 2. Normalize paragraph spacing & remove misplaced body page numbers / footer lines
        try:
            from docx.shared import Pt
            cleaning_sigs = [s.lower() for s in (layout_info.get("body_cleaning_signatures", []) if layout_info else []) if len(s.strip()) >= 3]
            for p in list(doc.paragraphs):
                # Extract ALL text inside paragraph including hyperlinks and nested runs
                full_text = "".join(p._element.xpath(".//w:t/text()")).strip()
                full_text_lower = full_text.lower()

                # Remove standalone body page numbers (e.g. '1', '2', '3') so they don't create blank pages
                if full_text.isdigit() and len(full_text) <= 3:
                    p._element.getparent().remove(p._element)
                    continue
                if re.match(r"^page\s*\d+(\s*of\s*\d+)?$", full_text, re.I):
                    p._element.getparent().remove(p._element)
                    continue

                # Remove misplaced footer signatures from body text
                if cleaning_sigs and any(sig in full_text_lower for sig in cleaning_sigs):
                    p._element.getparent().remove(p._element)
                    continue

                pf = p.paragraph_format
                if pf.space_before and pf.space_before.pt > 4:
                    pf.space_before = Pt(round(pf.space_before.pt * 0.4, 1))
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
        except Exception as spacing_err:
            logger.warning(f"Failed to normalize paragraph spacing: {spacing_err}")

        # 3. Image Z-Order & Layering:
        # Floating background watermark graphics (large / centered) -> behindDoc="1"
        # Content photos, signatures, stamps -> behindDoc="0"
        try:
            for anchor in doc._element.xpath(".//wp:anchor"):
                extent = anchor.xpath("wp:extent")
                cx = int(extent[0].get("cx", 0)) if extent else 0
                cy = int(extent[0].get("cy", 0)) if extent else 0
                # If image is very large (e.g. > 3 inches in both dimensions, cx >= 2,743,200 and cy >= 2,743,200) -> background emblem
                is_background_emblem = (cx >= 2743200 and cy >= 2743200)
                if is_background_emblem:
                    anchor.set("behindDoc", "1")
                else:
                    # Content photos (e.g. building photo, signature) stay in foreground
                    anchor.set("behindDoc", "0")
        except Exception as anchor_err:
            logger.warning(f"Failed to adjust drawing Z-order: {anchor_err}")

        # 4. Inject Modern DrawingML Watermark into Header if detected in source PDF
        watermark_info = layout_info.get("watermark") if layout_info else None
        if watermark_info and doc.sections:
            watermark_text = watermark_info.get("text", "Watermark")
            watermark_font = watermark_info.get("font", "DejaVu Sans")
            watermark_color = watermark_info.get("color", "2562EB")
            watermark_rot = watermark_info.get("rot", 19800000)
            wm_pages = watermark_info.get("pages_present", [0])

            drawingml_xml = (
                f'<w:p {nsdecls("w")} xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
                f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
                f'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">\n'
                f'  <w:pPr>\n'
                f'    <w:pStyle w:val="Header"/>\n'
                f'  </w:pPr>\n'
                f'  <w:r>\n'
                f'    <w:drawing>\n'
                f'      <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" '
                f'behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">\n'
                f'        <wp:simplePos x="0" y="0"/>\n'
                f'        <wp:positionH relativeFrom="page">\n'
                f'          <wp:align>center</wp:align>\n'
                f'        </wp:positionH>\n'
                f'        <wp:positionV relativeFrom="page">\n'
                f'          <wp:align>center</wp:align>\n'
                f'        </wp:positionV>\n'
                f'        <wp:extent cx="5486400" cy="1371600"/>\n'
                f'        <wp:effectExtent l="0" t="0" r="0" b="0"/>\n'
                f'        <wp:wrapNone/>\n'
                f'        <wp:docPr id="9999" name="WatermarkShape"/>\n'
                f'        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
                f'          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">\n'
                f'            <wps:wsp>\n'
                f'              <wps:cNvSpPr txBox="1"/>\n'
                f'              <wps:spPr>\n'
                f'                <a:xfrm rot="{watermark_rot}">\n'
                f'                  <a:off x="0" y="0"/>\n'
                f'                  <a:ext cx="5486400" cy="1371600"/>\n'
                f'                </a:xfrm>\n'
                f'                <a:prstGeom prst="rect">\n'
                f'                  <a:avLst/>\n'
                f'                </a:prstGeom>\n'
                f'                <a:noFill/>\n'
                f'                <a:ln>\n'
                f'                  <a:noFill/>\n'
                f'                </a:ln>\n'
                f'              </wps:spPr>\n'
                f'              <wps:txbx>\n'
                f'                <w:txbxContent>\n'
                f'                  <w:p>\n'
                f'                    <w:pPr>\n'
                f'                      <w:jc w:val="center"/>\n'
                f'                    </w:pPr>\n'
                f'                    <w:r>\n'
                f'                      <w:rPr>\n'
                f'                        <w:rFonts w:ascii="{watermark_font}" w:hAnsi="{watermark_font}" w:cs="{watermark_font}"/>\n'
                f'                        <w:b/>\n'
                f'                        <w:sz w:val="96"/>\n'
                f'                        <w:szCs w:val="96"/>\n'
                f'                        <w:color w:val="{watermark_color}"/>\n'
                f'                      </w:rPr>\n'
                f'                      <w:t>{watermark_text}</w:t>\n'
                f'                    </w:r>\n'
                f'                  </w:p>\n'
                f'                </w:txbxContent>\n'
                f'              </wps:txbx>\n'
                f'              <wps:bodyPr vert="horz" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr">\n'
                f'                <a:noFill/>\n'
                f'              </wps:bodyPr>\n'
                f'            </wps:wsp>\n'
                f'          </a:graphicData>\n'
                f'        </a:graphic>\n'
                f'      </wp:anchor>\n'
                f'    </w:drawing>\n'
                f'  </w:r>\n'
                f'</w:p>'
            )

            sec1 = doc.sections[0]
            if wm_pages == [0]:
                sec1.different_first_page_header_footer = True
                f_header = sec1.first_page_header
                for p in list(f_header.paragraphs):
                    p._element.getparent().remove(p._element)
                f_header._element.append(parse_xml(drawingml_xml))
            else:
                header = sec1.header
                for p in list(header.paragraphs):
                    p._element.getparent().remove(p._element)
                header._element.append(parse_xml(drawingml_xml))
                for sec in doc.sections[1:]:
                    sec.header.is_linked_to_previous = True

            logger.info(f"Injected modern DrawingML watermark '{watermark_text}' ({watermark_color}, rot={watermark_rot}) on pages {wm_pages}")

        # 5. Inject Extracted Footer / Dynamic Page Number into Word Footers
        has_page_numbers = layout_info.get("has_page_numbers", False) if layout_info else False
        has_text_footer = layout_info.get("has_text_footer", False) if layout_info else False
        footer_rows = layout_info.get("footer_rows", []) if layout_info else []

        if has_text_footer and footer_rows:
            # Inject structured 2-column contact/address footer with active hyperlink styling
            for sec in doc.sections:
                for ftr_target in (sec.first_page_footer, sec.footer):
                    for p in list(ftr_target.paragraphs):
                        p._element.getparent().remove(p._element)
                    for row in footer_rows:
                        left_t = row.get("left_text", "")
                        right_t = row.get("right_text", "")
                        f_name = row.get("font", "Calibri")
                        f_sz = row.get("size_half_pt", 16)
                        f_col = row.get("color", "71717A")

                        if right_t:
                            left_runs = render_rich_footer_runs(left_t, f_name, f_sz, f_col)
                            right_runs = render_rich_footer_runs(right_t, f_name, f_sz, f_col)
                            ftr_p_xml = (
                                f'<w:p {nsdecls("w")}>\n'
                                f'  <w:pPr>\n'
                                f'    <w:pStyle w:val="Footer"/>\n'
                                f'    <w:tabs>\n'
                                f'      <w:tab w:val="right" w:pos="9360"/>\n'
                                f'    </w:tabs>\n'
                                f'    <w:spacing w:before="0" w:after="0" w:line="220" w:lineRule="auto"/>\n'
                                f'  </w:pPr>\n'
                                f'{left_runs}\n'
                                f'  <w:r><w:tab/></w:r>\n'
                                f'{right_runs}\n'
                                f'</w:p>'
                            )
                        else:
                            left_runs = render_rich_footer_runs(left_t, f_name, f_sz, f_col)
                            ftr_p_xml = (
                                f'<w:p {nsdecls("w")}>\n'
                                f'  <w:pPr>\n'
                                f'    <w:pStyle w:val="Footer"/>\n'
                                f'    <w:spacing w:before="0" w:after="0" w:line="220" w:lineRule="auto"/>\n'
                                f'  </w:pPr>\n'
                                f'{left_runs}\n'
                                f'</w:p>'
                            )
                        ftr_target._element.append(parse_xml(ftr_p_xml))

        elif has_page_numbers:
            # Inject centered native Word page number field
            footer_page_xml = (
                f'<w:p {nsdecls("w")}>\n'
                f'  <w:pPr>\n'
                f'    <w:pStyle w:val="Footer"/>\n'
                f'    <w:jc w:val="center"/>\n'
                f'  </w:pPr>\n'
                f'  <w:fldSimple w:instr="PAGE"/>\n'
                f'</w:p>'
            )
            for sec in doc.sections:
                for ftr_target in (sec.first_page_footer, sec.footer):
                    for p in list(ftr_target.paragraphs):
                        p._element.getparent().remove(p._element)
                    ftr_target._element.append(parse_xml(footer_page_xml))

        doc.save(str(docx_path))
        return True
    except Exception as e:
        logger.warning(f"Failed in post-processing Word DOCX: {e}")
        return False


def parse_page_range(range_str: Optional[str], total_pages: int) -> List[int]:
    """
    Parses a user-supplied page range string (e.g. '1-3, 5, 8') into a sorted,
    deduplicated list of 0-indexed page indices.
    
    If range_str is empty or None, returns all pages [0, 1, ..., total_pages - 1].
    """
    if not range_str or not range_str.strip():
        return list(range(total_pages))

    indices = set()
    parts = range_str.split(",")

    for part in parts:
        clean_part = part.strip()
        if not clean_part:
            continue
        if "-" in clean_part:
            subparts = clean_part.split("-")
            if len(subparts) != 2:
                raise FileValidationException(
                    message=f"Invalid page range token: '{clean_part}'. Expected format 'X-Y'.",
                    code="INVALID_PAGE_RANGE",
                )
            try:
                start_p = int(subparts[0].strip())
                end_p = int(subparts[1].strip())
            except ValueError:
                raise FileValidationException(
                    message=f"Invalid page numbers in range: '{clean_part}'.",
                    code="INVALID_PAGE_RANGE",
                )
            if start_p < 1 or end_p < 1 or start_p > end_p:
                raise FileValidationException(
                    message=f"Invalid range bounds '{clean_part}': Start page must be >= 1 and <= End page.",
                    code="INVALID_PAGE_RANGE",
                )
            if end_p > total_pages:
                raise FileValidationException(
                    message=f"Page number {end_p} exceeds total document pages ({total_pages}).",
                    code="PAGE_OUT_OF_BOUNDS",
                    details={"total_pages": total_pages, "requested_page": end_p},
                )
            for p in range(start_p, end_p + 1):
                indices.add(p - 1)
        else:
            try:
                single_p = int(clean_part)
            except ValueError:
                raise FileValidationException(
                    message=f"Invalid page number token: '{clean_part}'.",
                    code="INVALID_PAGE_RANGE",
                )
            if single_p < 1 or single_p > total_pages:
                raise FileValidationException(
                    message=f"Page number {single_p} is out of bounds (document has {total_pages} pages).",
                    code="PAGE_OUT_OF_BOUNDS",
                    details={"total_pages": total_pages, "requested_page": single_p},
                )
            indices.add(single_p - 1)

    if not indices:
        raise FileValidationException(
            message="Parsed page range is empty.",
            code="EMPTY_PAGE_RANGE",
        )

    return sorted(list(indices))


def convert_pdf_sync(
    pdf_path: Path,
    docx_path: Path,
    pages: Optional[List[int]] = None,
    password: Optional[str] = None,
    progress_callback: Optional[Callable[[int, int, str, int], None]] = None,
) -> dict:
    """
    Synchronous conversion engine function.
    Converts PDF layout to DOCX format page-by-page using pdf2docx's native pipeline.
    Preserves original PDF styling, layout, graphics, and backgrounds with high fidelity.
    """
    start_time = time.time()
    unlocked_pdf_path: Optional[Path] = None
    effective_pdf_path = pdf_path

    # Step 1: If password is provided, decrypt to a temporary unencrypted PDF
    if password:
        try:
            doc = fitz.open(str(pdf_path))
            if doc.is_encrypted:
                auth_success = doc.authenticate(password)
                if auth_success == 0:
                    raise FileValidationException(
                        message="Incorrect password for encrypted PDF.",
                        code="INVALID_PASSWORD",
                    )
                unlocked_pdf_path = pdf_path.parent / f"unlocked_{pdf_path.name}"
                doc.save(str(unlocked_pdf_path))
                effective_pdf_path = unlocked_pdf_path
            doc.close()
        except Exception as e:
            if isinstance(e, FileValidationException):
                raise
            raise ConversionException(
                message=f"Failed to decrypt password-protected PDF: {e}",
                code="DECRYPTION_ERROR",
            )

    # Pre-scan PDF for layout structure, watermarks, images, and footers
    layout_info = analyze_pdf_layout_structure(effective_pdf_path)

    cv: Optional[Converter] = None
    try:
        cv = Converter(str(effective_pdf_path))
        doc_page_count = len(cv.fitz_doc)
        target_pages = pages if pages is not None else list(range(doc_page_count))
        total_target_count = len(target_pages)

        if total_target_count == 0:
            raise ConversionException(
                message="No pages selected for conversion.",
                code="NO_PAGES_SELECTED",
            )

        if progress_callback:
            progress_callback(0, total_target_count, "Initializing document layout analyzer", 10)

        if progress_callback:
            progress_callback(1, total_target_count, "Reconstructing vector layout & styles", 40)

        # Execute high-fidelity native conversion
        cv.convert(str(docx_path), pages=target_pages)

        if progress_callback:
            progress_callback(total_target_count, total_target_count, "Finalizing Word document formatting", 85)

        # Apply semantic post-processing
        if docx_path.exists():
            apply_word_post_processing(docx_path, layout_info=layout_info)

        duration = round(time.time() - start_time, 2)
        docx_size = docx_path.stat().st_size if docx_path.exists() else 0

        if progress_callback:
            progress_callback(
                total_target_count,
                total_target_count,
                "Conversion completed successfully",
                100,
            )

        logger.info(
            f"Successfully converted {total_target_count} page(s) to {docx_path.name} in {duration}s ({docx_size} bytes)"
        )

        return {
            "success": True,
            "pages_converted": total_target_count,
            "duration_seconds": duration,
            "docx_size_bytes": docx_size,
        }

    except Exception as e:
        logger.error(f"Conversion failed for {pdf_path.name}: {e}", exc_info=True)
        if isinstance(e, (FileValidationException, ConversionException)):
            raise
        raise ConversionException(
            message=f"Layout reconstruction failed: {str(e)}",
            code="CONVERSION_ERROR",
            details={"error_detail": str(e)},
        )
    finally:
        if cv:
            try:
                cv.close()
            except Exception:
                pass
        # Clean up unencrypted temporary PDF if one was created
        if unlocked_pdf_path and unlocked_pdf_path.exists():
            try:
                unlocked_pdf_path.unlink(missing_ok=True)
            except Exception:
                pass


async def convert_pdf_async(
    pdf_path: Path,
    docx_path: Path,
    pages: Optional[List[int]] = None,
    password: Optional[str] = None,
    progress_callback: Optional[Callable[[int, int, str, int], None]] = None,
    timeout_seconds: Optional[int] = None,
) -> dict:
    """
    Executes conversion asynchronously in a separate thread/process with a timeout watchdog.
    """
    timeout = timeout_seconds or settings.CONVERSION_TIMEOUT_SECONDS

    loop = asyncio.get_running_loop()
    try:
        return await asyncio.wait_for(
            loop.run_in_executor(
                None,
                convert_pdf_sync,
                pdf_path,
                docx_path,
                pages,
                password,
                progress_callback,
            ),
            timeout=timeout,
        )
    except asyncio.TimeoutError:
        raise ConversionException(
            message=f"Conversion timed out after {timeout} seconds.",
            code="CONVERSION_TIMEOUT",
            details={"timeout_seconds": timeout},
        )
