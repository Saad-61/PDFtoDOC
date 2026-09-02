import io
import json
import asyncio
from pathlib import Path
import pytest
import docx

from app.services.converter_engine import parse_page_range
from tests.test_validator_engine import create_sample_pdf, create_scanned_pdf


@pytest.mark.asyncio
async def test_api_validate_document(async_client, tmp_path):
    """Tests POST /api/v1/convert/validate preflight endpoint."""
    pdf_path = create_sample_pdf(tmp_path / "valid_test.pdf", num_pages=3)
    pdf_bytes = pdf_path.read_bytes()

    files = {"file": ("valid_test.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    response = await async_client.post("/api/v1/convert/validate", files=files)

    assert response.status_code == 200
    data = response.json()
    assert data["is_valid"] is True
    assert data["total_pages"] == 3
    assert data["is_encrypted"] is False
    assert data["is_scanned"] is False
    assert len(data["page_dimensions"]) == 3


@pytest.mark.asyncio
async def test_api_validate_encrypted_document_error(async_client, tmp_path):
    """Tests POST /api/v1/convert/validate with password-protected PDF without password."""
    pdf_path = create_sample_pdf(tmp_path / "locked.pdf", num_pages=2, password="mypassword")
    pdf_bytes = pdf_path.read_bytes()

    files = {"file": ("locked.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    response = await async_client.post("/api/v1/convert/validate", files=files)

    assert response.status_code == 401
    data = response.json()
    assert data["success"] is False
    assert data["error"]["code"] == "PASSWORD_REQUIRED"


@pytest.mark.asyncio
async def test_api_convert_stream(async_client, tmp_path):
    """Tests POST /api/v1/convert/stream direct conversion endpoint."""
    pdf_path = create_sample_pdf(tmp_path / "stream_test.pdf", num_pages=2)
    pdf_bytes = pdf_path.read_bytes()

    files = {"file": ("stream_test.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    response = await async_client.post("/api/v1/convert/stream", files=files)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "stream_test.docx" in response.headers["content-disposition"]

    docx_stream = io.BytesIO(response.content)
    doc = docx.Document(docx_stream)
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert "Sample Document" in full_text


@pytest.mark.asyncio
async def test_api_convert_stream_page_range(async_client, tmp_path):
    """Tests POST /api/v1/convert/stream with selective page range."""
    pdf_path = create_sample_pdf(tmp_path / "range_test.pdf", num_pages=3)
    pdf_bytes = pdf_path.read_bytes()

    files = {"file": ("range_test.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    data = {"page_range": "2"}
    response = await async_client.post("/api/v1/convert/stream", files=files, data=data)

    assert response.status_code == 200
    doc = docx.Document(io.BytesIO(response.content))
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert "Page 2" in full_text
    assert "Page 1" not in full_text


@pytest.mark.asyncio
async def test_api_async_job_and_sse_events(async_client, tmp_path):
    """Tests complete async job lifecycle: create job, stream SSE events, query status, download docx."""
    pdf_path = create_sample_pdf(tmp_path / "async_api_test.pdf", num_pages=2)
    pdf_bytes = pdf_path.read_bytes()

    # 1. Enqueue Job
    files = {"file": ("async_api_test.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    create_res = await async_client.post("/api/v1/convert/jobs", files=files)

    assert create_res.status_code == 202
    job_data = create_res.json()
    job_id = job_data["job_id"]
    assert job_data["status"] == "QUEUED"
    assert f"/api/v1/jobs/{job_id}/events" == job_data["events_url"]

    # 2. Wait a moment for background worker to complete
    await asyncio.sleep(1.5)

    # 3. Query Job Status
    status_res = await async_client.get(f"/api/v1/jobs/{job_id}")
    assert status_res.status_code == 200
    status_data = status_res.json()
    assert status_data["status"] == "COMPLETED"
    assert status_data["percent"] == 100
    assert status_data["download_url"] is not None

    # 4. Download Converted DOCX
    download_res = await async_client.get(f"/api/v1/jobs/{job_id}/download")
    assert download_res.status_code == 200
    assert download_res.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    doc = docx.Document(io.BytesIO(download_res.content))
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert "Sample Document" in full_text


@pytest.mark.asyncio
async def test_api_unicode_filename_header(async_client, tmp_path):
    """Tests RFC 5987 Unicode filename preservation in Content-Disposition."""
    pdf_path = create_sample_pdf(tmp_path / "résumé_2026.pdf", num_pages=1)
    pdf_bytes = pdf_path.read_bytes()

    files = {"file": ("résumé_2026.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
    res = await async_client.post("/api/v1/convert/stream", files=files)

    assert res.status_code == 200
    cd_header = res.headers["content-disposition"]
    assert "filename*=UTF-8''" in cd_header
    assert "r%C3%A9sum%C3%A9_2026.docx" in cd_header


@pytest.mark.asyncio
async def test_api_not_found_job(async_client):
    """Tests querying non-existent job ID."""
    res = await async_client.get("/api/v1/jobs/non-existent-uuid-123")
    assert res.status_code == 404
    assert res.json()["error"]["code"] == "RESOURCE_NOT_FOUND"


@pytest.mark.asyncio
async def test_api_real_resume_conversion(async_client):
    """Tests end-to-end conversion on the real SaadAsifResume.pdf in the workspace."""
    resume_file = Path("SaadAsifResume.pdf")
    if not resume_file.exists():
        resume_file = Path(__file__).resolve().parent.parent.parent / "SaadAsifResume.pdf"
    
    if resume_file.exists():
        files = {"file": ("SaadAsifResume.pdf", open(resume_file, "rb"), "application/pdf")}
        res = await async_client.post("/api/v1/convert/stream", files=files)
        assert res.status_code == 200
        assert res.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = docx.Document(io.BytesIO(res.content))
        assert len(doc.paragraphs) > 0

