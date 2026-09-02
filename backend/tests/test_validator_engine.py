import io
from pathlib import Path
import pytest
import fitz
import docx
from PIL import Image

from app.core.exceptions import FileValidationException, PasswordRequiredException, ConversionException
from app.services.pdf_validator import pdf_validator
from app.services.converter_engine import parse_page_range, convert_pdf_sync, convert_pdf_async
from app.services.job_manager import job_manager, JobStatus


def create_sample_pdf(file_path: Path, num_pages: int = 3, password: str = None) -> Path:
    """Helper to generate a valid synthetic PDF with text and geometry."""
    doc = fitz.open()
    for i in range(num_pages):
        page = doc.new_page(width=595, height=842)  # A4 size
        # Add heading and body text
        page.insert_text((50, 60), f"Sample Document Page {i + 1}", fontsize=18)
        page.insert_text((50, 100), "This is a paragraph of text used for testing layout reconstruction in Word.", fontsize=11)
        # Add a simple rectangle shape
        page.draw_rect(fitz.Rect(50, 130, 300, 180), color=(0.2, 0.4, 0.8), width=1)
        page.insert_text((60, 155), f"Table / Box Item {i + 1}", fontsize=10)

    if password:
        # Encrypt with owner and user password
        perm = int(
            fitz.PDF_PERM_ACCESSIBILITY
            | fitz.PDF_PERM_PRINT
            | fitz.PDF_PERM_COPY
            | fitz.PDF_PERM_ANNOTATE
        )
        doc.save(
            str(file_path),
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw="owner123",
            user_pw=password,
            permissions=perm,
        )
    else:
        doc.save(str(file_path))
    doc.close()
    return file_path


def create_scanned_pdf(file_path: Path) -> Path:
    """Helper to generate a synthetic pure-raster scanned PDF (image only, no text)."""
    # Create a small dummy image
    img = Image.new("RGB", (300, 400), color=(240, 240, 240))
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format="PNG")
    img_bytes = img_byte_arr.getvalue()

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(fitz.Rect(50, 50, 400, 500), stream=img_bytes)
    doc.save(str(file_path))
    doc.close()
    return file_path


def test_validator_valid_pdf(tmp_path):
    """Tests that a standard valid PDF passes validation with correct page count and dimensions."""
    pdf_path = create_sample_pdf(tmp_path / "valid.pdf", num_pages=3)
    result = pdf_validator.validate(pdf_path)

    assert result.is_valid is True
    assert result.total_pages == 3
    assert result.is_encrypted is False
    assert result.is_scanned is False
    assert len(result.page_dimensions) == 3
    assert result.page_dimensions[0].width == 595.0
    assert result.page_dimensions[0].height == 842.0


def test_validator_password_protected(tmp_path):
    """Tests password detection and validation for encrypted PDFs."""
    pdf_path = create_sample_pdf(tmp_path / "protected.pdf", num_pages=2, password="secretpassword")

    # 1. Without password, should raise PasswordRequiredException
    with pytest.raises(PasswordRequiredException) as exc_info:
        pdf_validator.validate(pdf_path)
    assert exc_info.value.code == "PASSWORD_REQUIRED"

    # 2. With incorrect password, should raise FileValidationException
    with pytest.raises(FileValidationException) as exc_info:
        pdf_validator.validate(pdf_path, password="wrongpassword")
    assert exc_info.value.code == "INVALID_PASSWORD"

    # 3. With correct password, should succeed
    result = pdf_validator.validate(pdf_path, password="secretpassword")
    assert result.is_valid is True
    assert result.is_encrypted is True
    assert result.total_pages == 2


def test_validator_scanned_pdf(tmp_path):
    """Tests detection heuristic for scanned raster PDFs."""
    pdf_path = create_scanned_pdf(tmp_path / "scanned.pdf")
    result = pdf_validator.validate(pdf_path)
    assert result.is_valid is True
    assert result.is_scanned is True


def test_validator_corrupt_file(tmp_path):
    """Tests that files without valid magic bytes or malformed bytes fail validation."""
    bad_file = tmp_path / "corrupt.pdf"
    bad_file.write_bytes(b"NOT_A_PDF_DATA_HEADER_HERE")

    with pytest.raises(FileValidationException) as exc_info:
        pdf_validator.validate(bad_file)
    assert exc_info.value.code == "INVALID_PDF_HEADER"


def test_page_range_parser():
    """Tests page range parsing logic."""
    assert parse_page_range(None, 5) == [0, 1, 2, 3, 4]
    assert parse_page_range("", 5) == [0, 1, 2, 3, 4]
    assert parse_page_range("1-3, 5", 5) == [0, 1, 2, 4]
    assert parse_page_range("2", 5) == [1]
    assert parse_page_range("3-3", 5) == [2]

    # Out of bounds
    with pytest.raises(FileValidationException) as exc:
        parse_page_range("1-6", 5)
    assert exc.value.code == "PAGE_OUT_OF_BOUNDS"

    # Invalid range syntax
    with pytest.raises(FileValidationException) as exc:
        parse_page_range("5-2", 5)
    assert exc.value.code == "INVALID_PAGE_RANGE"


def test_convert_pdf_sync_full_document(tmp_path):
    """Tests synchronous conversion of full PDF to valid DOCX document."""
    pdf_path = create_sample_pdf(tmp_path / "sample_convert.pdf", num_pages=2)
    docx_path = tmp_path / "converted.docx"

    progress_events = []
    def on_progress(page, total, stage, percent):
        progress_events.append((page, total, stage, percent))

    result = convert_pdf_sync(
        pdf_path=pdf_path,
        docx_path=docx_path,
        progress_callback=on_progress,
    )

    assert result["success"] is True
    assert result["pages_converted"] == 2
    assert docx_path.exists()
    assert docx_path.stat().st_size > 0
    assert len(progress_events) >= 3

    # Verify DOCX is valid and readable by python-docx
    doc = docx.Document(str(docx_path))
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert "Sample Document" in full_text


def test_convert_pdf_sync_page_range(tmp_path):
    """Tests converting a specific page subset (e.g. only page 2)."""
    pdf_path = create_sample_pdf(tmp_path / "sample_range.pdf", num_pages=3)
    docx_path = tmp_path / "converted_page2.docx"

    target_pages = parse_page_range("2", 3)  # [1]
    result = convert_pdf_sync(
        pdf_path=pdf_path,
        docx_path=docx_path,
        pages=target_pages,
    )

    assert result["success"] is True
    assert result["pages_converted"] == 1
    assert docx_path.exists()

    doc = docx.Document(str(docx_path))
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert "Page 2" in full_text
    assert "Page 1" not in full_text


@pytest.mark.asyncio
async def test_convert_pdf_async(tmp_path):
    """Tests asynchronous non-blocking conversion wrapper."""
    pdf_path = create_sample_pdf(tmp_path / "async_sample.pdf", num_pages=2)
    docx_path = tmp_path / "async_converted.docx"

    result = await convert_pdf_async(
        pdf_path=pdf_path,
        docx_path=docx_path,
    )

    assert result["success"] is True
    assert docx_path.exists()


@pytest.mark.asyncio
async def test_job_manager_workflow(tmp_path):
    """Tests JobManager state management and event broadcast subscriptions."""
    job_id = "job-test-uuid-456"
    session_id = "session-123"
    filename = "report.pdf"

    # 1. Create job
    record = await job_manager.create_job(job_id, session_id, filename, total_pages=5)
    assert record.job_id == job_id
    assert record.status == JobStatus.QUEUED

    # 2. Subscribe listener
    queue = await job_manager.subscribe(job_id)
    initial_event = await queue.get()
    assert initial_event["status"] == "QUEUED"

    # 3. Update progress
    await job_manager.update_progress(job_id, current_page=2, total_pages=5, stage="Reconstructing page 2", percent=40)
    progress_event = await queue.get()
    assert progress_event["status"] == "PROCESSING"
    assert progress_event["current_page"] == 2
    assert progress_event["percent"] == 40

    # 4. Complete job
    dummy_docx = tmp_path / "output.docx"
    dummy_docx.write_bytes(b"PK...docx")
    await job_manager.complete_job(job_id, docx_path=dummy_docx, duration_seconds=1.85)

    completed_event = await queue.get()
    assert completed_event["status"] == "COMPLETED"
    assert completed_event["percent"] == 100
    assert completed_event["duration_seconds"] == 1.85

    # 5. Unsubscribe
    await job_manager.unsubscribe(job_id, queue)
